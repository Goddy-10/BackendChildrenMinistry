from flask import Blueprint, request, jsonify,send_file

from flask_jwt_extended import jwt_required,get_jwt_identity
from datetime import datetime
from app.models import  FinanceEntry, Project, Mission, Department, NewMember,Expenditure,MissionPartner
from app.extensions import db
from sqlalchemy import and_
from io import BytesIO
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas


adults_bp = Blueprint("adults_bp", __name__, url_prefix="/adults")








# ------------ NEW PDF + DOCX EXPORT (supports Income + Expenditure + Net Balance) ------------



@adults_bp.route("/finance/export/pdf", methods=["GET"])
def export_finance_pdf():
    start_str = request.args.get("start")
    end_str = request.args.get("end")

    start_date = datetime.strptime(start_str, "%Y-%m-%d").date() if start_str else None
    end_date = datetime.strptime(end_str, "%Y-%m-%d").date() if end_str else None

    # Fetch both tables
    income_query = FinanceEntry.query.order_by(FinanceEntry.date)
    exp_query = Expenditure.query.order_by(Expenditure.date)

    if start_date:
        income_query = income_query.filter(FinanceEntry.date >= start_date)
        exp_query = exp_query.filter(Expenditure.date >= start_date)
    if end_date:
        income_query = income_query.filter(FinanceEntry.date <= end_date)
        exp_query = exp_query.filter(Expenditure.date <= end_date)

    incomes = income_query.all()
    expenditures = exp_query.all()

    # Calculations
    total_main = sum(float(e.main_church or 0) for e in incomes)
    total_children = sum(float(e.children_ministry or 0) for e in incomes)
    total_income = total_main + total_children
    total_expense = sum(float(e.amount or 0) for e in expenditures)
    net_balance = total_income - total_expense

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    margin = 1.5 * cm
    y = height - margin

    p.setFont("Helvetica-Bold", 16)
    p.drawString(margin, y, "CHURCH FINANCE REPORT")
    y -= 30

    p.setFont("Helvetica", 11)
    period = "All Time"
    if start_date and end_date:
        period = f"{start_date} to {end_date}"
    elif start_date:
        period = f"From {start_date}"
    elif end_date:
        period = f"Up to {end_date}"
    p.drawString(margin, y, f"Period: {period}")
    y -= 40

    # Summary Box
    p.setFont("Helvetica-Bold", 12)
    p.drawString(margin, y, "SUMMARY")
    y -= 20
    p.setFont("Helvetica", 11)
    p.drawString(margin + 20, y, f"Main Church Income:        KSh {total_main:,.2f}")
    y -= 18
    p.drawString(margin + 20, y, f"Children Ministry Income:  KSh {total_children:,.2f}")
    y -= 18
    p.drawString(margin + 20, y, f"Total Income:              KSh {total_income:,.2f}")
    y -= 18
    p.drawString(margin + 20, y, f"Total Expenditure:         KSh {total_expense:,.2f}")
    y -= 25
    p.setFont("Helvetica-Bold", 14)
    balance_color = "black" if net_balance >= 0 else "red"
    p.setFillColor(balance_color)
    p.drawString(margin + 20, y, f"NET BALANCE:               KSh {net_balance:,.2f}")
    p.setFillColor("black")
    y -= 50

    # Table Header
    p.setFont("Helvetica-Bold", 10)
    headers = ["Date", "Type", "Description", "Income", "Expense"]
    x_positions = [margin, margin + 2*cm, margin + 5*cm, margin + 11*cm, margin + 14*cm]
    for text, x in zip(headers, x_positions):
        p.drawString(x, y, text)
    y -= 15
    p.line(margin, y, width - margin, y)
    y -= 10

    # Data rows
    p.setFont("Helvetica", 9)
    all_entries = (
        [(e.date, "Income", f"{e.service_type or 'Offering'} (Main)" if e.main_church else f"{e.service_type or 'Offering'} (Children)", 
          float(e.main_church or e.children_ministry), 0) for e in incomes] +
        [(e.date, "Expenditure", e.details or "No details", 0, float(e.amount)) for e in expenditures]
    )
    all_entries.sort(key=lambda x: x[0], reverse=True)

    for date, typ, desc, inc, exp in all_entries:
        if y < 100:
            p.showPage()
            y = height - margin
        p.drawString(x_positions[0], y, date.strftime("%Y-%m-%d"))
        p.drawString(x_positions[1], y, typ)
        p.drawString(x_positions[2], y, desc[:40])
        if inc > 0:
            p.drawRightString(x_positions[3] + 1*cm, y, f"{inc:,.2f}")
        if exp > 0:
            p.drawRightString(x_positions[4] + 1*cm, y, f"{exp:,.2f}")
        y -= 18

    p.showPage()
    p.save()
    buffer.seek(0)

    filename = f"finance_report_{start_str or 'all'}_to_{end_str or 'now'}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype="application/pdf")



#------------WORD--------#

@adults_bp.route("/finance/export/docx", methods=["GET"])
def export_finance_docx():
    start_str = request.args.get("start")
    end_str = request.args.get("end")

    start_date = datetime.strptime(start_str, "%Y-%m-%d").date() if start_str else None
    end_date = datetime.strptime(end_str, "%Y-%m-%d").date() if end_str else None

    income_query = FinanceEntry.query.order_by(FinanceEntry.date)
    exp_query = Expenditure.query.order_by(Expenditure.date)

    if start_date:
        income_query = income_query.filter(FinanceEntry.date >= start_date)
        exp_query = exp_query.filter(Expenditure.date >= start_date)
    if end_date:
        income_query = income_query.filter(FinanceEntry.date <= end_date)
        exp_query = exp_query.filter(Expenditure.date <= end_date)

    incomes = income_query.all()
    expenditures = exp_query.all()

    total_main = sum(float(e.main_church or 0) for e in incomes)
    total_children = sum(float(e.children_ministry or 0) for e in incomes)
    total_income = total_main + total_children
    total_expense = sum(float(e.amount or 0) for e in expenditures)
    net_balance = total_income - total_expense

    doc = Document()
    doc.add_heading("CHURCH FINANCE REPORT", 0)

    p = doc.add_paragraph()
    p.add_run("Period: ").bold = True
    p.add_run(f"{start_date or 'Beginning'} → {end_date or 'Present'}")

    doc.add_paragraph()  # spacer

    # Summary Table
    table = doc.add_table(rows=1, cols=2, style="Table Grid")
    hdr = table.rows[0].cells
    hdr[0].text = "Description"
    hdr[1].text = "Amount (KSh)"

    summary_data = [
        ("Main Church Income", f"{total_main:,.2f}"),
        ("Children Ministry Income", f"{total_children:,.2f}"),
        ("Total Income", f"{total_income:,.2f}"),
        ("Total Expenditure", f"{total_expense:,.2f}"),
        ("NET BALANCE", f"{net_balance:,.2f}"),
    ]
    for desc, amt in summary_data:
        row = table.add_row().cells
        row[0].text = desc
        row[1].text = amt
        if "NET" in desc:
            row[0].paragraphs[0].runs[0].bold = True
            row[1].paragraphs[0].runs[0].bold = True
            if net_balance < 0:
                row[1].paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(200, 0, 0)

    doc.add_page_break()

    # Detailed Transactions
    doc.add_heading("Detailed Transactions", level=1)
    table = doc.add_table(rows=1, cols=5, style="Table Grid")
    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Type"
    hdr[2].text = "Description"
    hdr[3].text = "Income"
    hdr[4].text = "Expense"

    all_entries = (
        [(e.date, "Income", e.service_type or "Offering", float(e.main_church or e.children_ministry), 0) for e in incomes] +
        [(e.date, "Expenditure", e.details or "", 0, float(e.amount)) for e in expenditures]
    )
    all_entries.sort(key=lambda x: x[0], reverse=True)

    for date, typ, desc, inc, exp in all_entries:
        row = table.add_row().cells
        row[0].text = date.strftime("%Y-%m-%d")
        row[1].text = typ
        row[2].text = desc
        row[3].text = f"{inc:,.2f}" if inc > 0 else ""
        row[4].text = f"{exp:,.2f}" if exp > 0 else ""

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"finance_report_{start_str or 'all'}_to_{end_str or 'now'}.docx"
    return send_file(buffer, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# -------------------- FINANCE ENTRIES --------------------


# ────────────── GET ALL ENTRIES (income + expenditure) ──────────────
@adults_bp.route("/finance/all", methods=["GET","OPTIONS"])
@jwt_required()
def get_all_finances():

    income = [e.to_dict() for e in FinanceEntry.query.order_by(FinanceEntry.date.desc()).all()]
    expenses = [e.to_dict() for e in Expenditure.query.order_by(Expenditure.date.desc()).all()]
    # Combine and sort by date descending
    all_entries = sorted(income + expenses, key=lambda x: x["date"], reverse=True)
    return jsonify(all_entries), 200

# ────────────── ADD INCOME ──────────────
@adults_bp.route("/finance/income", methods=["POST"])
@jwt_required()
def add_income():
    current_user_id = get_jwt_identity()  # ← this pulls user ID from JWT token
    data = request.get_json()
    entry = FinanceEntry(
        date=datetime.strptime(data["date"], "%Y-%m-%d"),
        service_type=data.get("service_type"),
        main_church=data.get("main_church", 0),
        children_ministry=data.get("children_ministry", 0),
        created_by=current_user_id
    )
    db.session.add(entry)
    db.session.commit()
    db.session.refresh(entry)
    return jsonify(entry.to_dict()), 201

# ─────────────── EXPENDITURE ──────────────
@adults_bp.route("/finance/expenditure", methods=["POST"])
@jwt_required()
def add_expenditure():
    current_user_id = get_jwt_identity()  # ← this pulls user ID from JWT token
    data = request.get_json()
    exp = Expenditure(
        date=datetime.strptime(data["date"], "%Y-%m-%d"),
        amount=data["amount"],
        details=data["details"],
        created_by=current_user_id
    )
    db.session.add(exp)
    db.session.commit()
    db.session.refresh(exp)
    return jsonify(exp.to_dict()), 201

# ────────────── UNIFIED UPDATE (works for both tables) ──────────────
@adults_bp.route("/finance/entry/<int:id>", methods=["PATCH"])
@jwt_required()
def update_entry(id):\

    data = request.get_json()
    current_user_id = get_jwt_identity()  # ← this pulls user ID from JWT token

    # Try income table first
    entry = FinanceEntry.query.get(id)
    if not entry:
        entry = Expenditure.query.get_or_404(id)  # will 404 if not found

    if isinstance(entry, FinanceEntry):
        if "date" in data:
            entry.date = datetime.strptime(data["date"], "%Y-%m-%d")
        if "service_type" in data:
            entry.service_type = data["service_type"]
        if "main_church" in data:
            entry.main_church = data["main_church"]
        if "children_ministry" in data:
            entry.children_ministry = data["children_ministry"]

    else:  # Expenditure
        if "date" in data:
            entry.date = datetime.strptime(data["date"], "%Y-%m-%d")
        if "amount" in data:
            entry.amount = data["amount"]
        if "details" in data:
            entry.details = data["details"]

    db.session.commit()
    return jsonify(entry.to_dict()), 200

# ────────────── UNIFIED DELETE ──────────────
@adults_bp.route("/finance/entry/<int:id>", methods=["DELETE"])
@jwt_required()
def delete_finance_entry(id):
    entry = FinanceEntry.query.get(id)
    current_user_id = get_jwt_identity()  # ← this pulls user ID from JWT token
    if not entry:
        entry = Expenditure.query.get(id)
        if not entry:
            return jsonify({"error": "Not found"}), 404

    db.session.delete(entry)
    db.session.commit()
    return jsonify({"message": "Deleted"}), 200


# -------------------- PROJECTS --------------------
@adults_bp.route("/projects", methods=["GET"])
def get_projects():
    projects = Project.query.all()
    return jsonify([p.to_dict() for p in projects]), 200

@adults_bp.route("/projects", methods=["POST"])
def add_project():
    data = request.get_json()

    # convert date strings to date objects if provided
    start_date = data.get("start_date")
    end_date = data.get("end_date")
    if start_date:
        start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
    if end_date:
        end_date = datetime.strptime(end_date, "%Y-%m-%d").date()

    project = Project(
        title=data["title"],
        description=data.get("description"),
        status=data.get("status", "planned"),
        start_date=start_date,
        end_date=end_date,
        created_by=data.get("created_by")
    )
    db.session.add(project)
    db.session.commit()
    return jsonify(project.to_dict()), 201

@adults_bp.route("/projects/<int:id>", methods=["PATCH"])
def update_project(id):
    project = Project.query.get_or_404(id)
    data = request.get_json()

    for field in ["title", "description", "status", "start_date", "end_date"]:
        if field in data:
            value = data[field]
            if field in ["start_date", "end_date"] and value:
                value = datetime.strptime(value, "%Y-%m-%d").date()
            setattr(project, field, value)

    db.session.commit()
    return jsonify(project.to_dict()), 200

@adults_bp.route("/projects/<int:id>", methods=["DELETE"])
def delete_project(id):
    project = Project.query.get_or_404(id)
    db.session.delete(project)
    db.session.commit()
    return jsonify({"message": "Project deleted"}), 200


# -------------------- MISSIONS --------------------
@adults_bp.route("/missions", methods=["GET"])
def get_missions():
    missions = Mission.query.order_by(Mission.date.desc()).all()
    return jsonify([m.to_dict() for m in missions]), 200




@adults_bp.route("/missions", methods=["POST"])
def add_mission():
    data = request.get_json()

    mission = Mission(
        title=data["title"],
        date=datetime.strptime(data["date"], "%Y-%m-%d").date(),
        location=data["location"],
        souls_won=data["souls_won"]
    )

    db.session.add(mission)
    db.session.commit()

    return jsonify(mission.to_dict()), 201




@adults_bp.route("/missions/<int:id>", methods=["PATCH"])
def update_mission(id):
    mission = Mission.query.get_or_404(id)
    data = request.get_json()

    if "title" in data:
        mission.title = data["title"]
    if "location" in data:
        mission.location = data["location"]
    if "date" in data:
        mission.date = datetime.strptime(data["date"], "%Y-%m-%d").date()
    if "souls_won" in data:
        mission.souls_won=data["souls_won"]   

    db.session.commit()
    return jsonify(mission.to_dict()), 200




@adults_bp.route("/missions/<int:id>", methods=["DELETE"])
def delete_mission(id):
    mission = Mission.query.get_or_404(id)
    db.session.delete(mission)
    db.session.commit()
    return jsonify({"message": "Mission deleted"}), 200





#-----MISSION PARTNERS ROUTES--------#


@adults_bp.route("/missions/<int:mission_id>/partners", methods=["GET"])
def get_mission_partners(mission_id):
    Mission.query.get_or_404(mission_id)  # ensure mission exists
    partners = MissionPartner.query.filter_by(mission_id=mission_id).all()
    return jsonify([p.to_dict() for p in partners]), 200


@adults_bp.route("/missions/<int:mission_id>/partners", methods=["POST"])
def add_mission_partner(mission_id):
    Mission.query.get_or_404(mission_id)
    data = request.get_json()

    partner = MissionPartner(
        mission_id=mission_id,
        name=data.get("partner_name"),
        support=data.get("support"),
        contact=data.get("contact")
    )

    db.session.add(partner)
    db.session.commit()

    return jsonify(partner.to_dict()), 201


@adults_bp.route("/missions/partners/<int:id>", methods=["PATCH"])
def update_mission_partner(id):
    partner = MissionPartner.query.get_or_404(id)
    data = request.get_json()

    if "name" in data:
        partner.name = data["name"]
    if "support" in data:
        partner.support = data["support"]
    if "contact" in data:
        partner.contact = data["contact"]

    db.session.commit()
    return jsonify(partner.to_dict()), 200



@adults_bp.route("/missions/partners/<int:id>", methods=["DELETE"])
def delete_mission_partner(id):
    partner = MissionPartner.query.get_or_404(id)
    db.session.delete(partner)
    db.session.commit()
    return jsonify({"message": "Partner removed"}), 200


# -------------------- DEPARTMENTS --------------------
@adults_bp.route("/departments", methods=["GET"])
def get_departments():
    depts = Department.query.all()
    return jsonify([d.to_dict() for d in depts]), 200

@adults_bp.route("/departments", methods=["POST"])
def add_department():
    data = request.get_json()
    dept = Department(
        name=data["name"],
        description=data.get("description"),
        contact_person=data.get("contact_person"),
        contact_phone=data.get("contact_phone"),
        contact_email=data.get("contact_email"),
    )
    db.session.add(dept)
    db.session.commit()
    return jsonify(dept.to_dict()), 201

@adults_bp.route("/departments/<int:id>", methods=["PATCH"])
def update_department(id):
    dept = Department.query.get_or_404(id)
    data = request.get_json()
    for field in ["name", "description", "contact_person", "contact_phone", "contact_email"]:
        if field in data:
            setattr(dept, field, data[field])
    db.session.commit()
    return jsonify(dept.to_dict()), 200

@adults_bp.route("/departments/<int:id>", methods=["DELETE"])
def delete_department(id):
    dept = Department.query.get_or_404(id)
    db.session.delete(dept)
    db.session.commit()
    return jsonify({"message": "Department deleted"}), 200


# -------------------- NEW MEMBERS --------------------
@adults_bp.route("/new-members", methods=["GET"])
def get_new_members():
    members = NewMember.query.all()
    return jsonify([m.to_dict() for m in members]), 200

@adults_bp.route("/new-members", methods=["POST"])
def add_new_member():
    data = request.get_json()
    member = NewMember(
        name=data["name"],
        phone=data.get("phone"),
        email=data.get("email"),
        join_date=datetime.strptime(data["join_date"], "%Y-%m-%d") if data.get("join_date") else datetime.utcnow(),
        notes=data.get("notes"),
        sunday_class_id=data.get("sunday_class_id"),
        department_id=data.get("department_id"),
    )
    db.session.add(member)
    db.session.commit()
    return jsonify(member.to_dict()), 201

@adults_bp.route("/new-members/<int:id>", methods=["PATCH"])
def update_new_member(id):
    member = NewMember.query.get_or_404(id)
    data = request.get_json()
    for field in ["name", "phone", "email", "notes", "sunday_class_id", "department_id"]:
        if field in data:
            setattr(member, field, data[field])

    # Handle join_date separately
    if "join_date" in data:
        member.join_date = datetime.strptime(data["join_date"], "%Y-%m-%d") if data["join_date"] else None

    db.session.commit()
    return jsonify(member.to_dict()), 200

@adults_bp.route("/new-members/<int:id>", methods=["DELETE"])
def delete_new_member(id):
    member = NewMember.query.get_or_404(id)
    db.session.delete(member)
    db.session.commit()
    return jsonify({"message": "New member deleted"}), 200